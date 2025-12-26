"""
Create comprehensive export tool that captures ALL agent dialogue
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Enhanced export tool that accumulates all agent outputs
NEW_SOURCE_CODE = '''from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Global storage for all agent outputs
_ALL_OUTPUTS = []
_EXPORT_COUNT = 0

def _set_rtl(para):
    """Set RTL direction for Arabic text"""
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

# Agent name mapping
AGENTS = {
    "series_bible": "وكيل قفل الشخصيات (Series Bible Lock)",
    "research": "وكيل البحث العميق (Deep Research)",
    "curriculum": "مهندس المنهج (Curriculum Engineer)",
    "planner": "وكيل تخطيط الحلقات (Episode Planner)",
    "script_v1": "كاتب السيناريو v1 (Scenarist)",
    "sheikh": "المراجع الشرعي (Sheikh)",
    "education": "المستشار التربوي (Education)",
    "psychiatrist": "الطبيب النفسي (Psychiatrist)",
    "bible_compliance": "وكيل اتساق السلسلة (Series Bible Agent)",
    "youtube": "مستشار يوتيوب (YouTube)",
    "revision": "وكيل تلخيص المراجعات (Revision Digest)",
    "script_v2": "كاتب السيناريو v2 (Rewrite)",
    "coverage": "وكيل تقرير التغطية (Coverage Report)",
    "language": "المدقق اللغوي (Language Consultant)",
    "exporter": "وكيل التصدير (Exporter)",
    "moderator": "وكيل المدير التنفيذي (Moderator)"
}

def _detect_agent(txt):
    """Detect which agent produced this content"""
    t = txt.lower()
    if "character_lock" in t or "location_lock" in t:
        return AGENTS["series_bible"]
    elif "child_friendly_definition" in t or "story_seeds" in t:
        return AGENTS["research"]
    elif "learning_objective" in t and "key_takeaways" in t:
        return AGENTS["curriculum"]
    elif "beat_sheet" in t or "hook" in t and "closing_message" in t:
        return AGENTS["planner"]
    elif "script_v1_title" in t or "script_v1_text" in t:
        return AGENTS["script_v1"]
    elif "script_v2_title" in t or "script_v2_text" in t:
        return AGENTS["script_v2"]
    elif "script_text_final" in t and "top_fixes" in t:
        return AGENTS["language"]
    elif "patch_plan" in t and "must_fix_top5" in t:
        return AGENTS["revision"]
    elif "changelog" in t and "risk_check" in t:
        return AGENTS["coverage"]
    elif "bible_compliance" in t:
        return AGENTS["bible_compliance"]
    elif "policy_and_safety" in t or "title_and_thumbnail" in t:
        return AGENTS["youtube"]
    elif "must_fix" in t or "should_fix" in t:
        if "شرع" in txt or "فقه" in txt:
            return AGENTS["sheikh"]
        elif "تربو" in txt:
            return AGENTS["education"]
        elif "نفس" in txt or "عاطف" in txt:
            return AGENTS["psychiatrist"]
    elif "export_status" in t:
        return AGENTS["exporter"]
    return AGENTS["moderator"]

def export_to_word(content: str, title: str = "") -> dict:
    """
    Accumulate agent outputs and create comprehensive Word document.
    Each call adds content; the document is rebuilt with ALL content.
    """
    global _ALL_OUTPUTS, _EXPORT_COUNT
    
    try:
        folder = "C:\\\\Users\\\\srab1.SAMEH-NVME\\\\Downloads\\\\AutoGen Studio Final\\\\AutoGen Studio\\\\AutoGen Studio\\\\Script"
        os.makedirs(folder, exist_ok=True)
        
        # Detect agent
        agent_name = _detect_agent(str(content))
        
        # Add to outputs
        _EXPORT_COUNT += 1
        _ALL_OUTPUTS.append({
            "order": _EXPORT_COUNT,
            "agent": agent_name,
            "content": str(content),
            "time": datetime.now().strftime("%H:%M:%S")
        })
        
        # Create document with ALL outputs
        doc = Document()
        
        # Title page (Arabic)
        ttl = doc.add_heading("", level=0)
        run = ttl.add_run("تقرير حوار وكلاء رفيق الشامل")
        run.font.size = Pt(24)
        _set_rtl(ttl)
        ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sub = doc.add_paragraph()
        sub.add_run("Rafeeq Complete Workflow Dialogue")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = doc.add_paragraph()
        info.add_run(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        _set_rtl(info)
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cnt = doc.add_paragraph()
        cnt.add_run(f"عدد المخرجات المسجلة: {len(_ALL_OUTPUTS)}")
        _set_rtl(cnt)
        cnt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of contents
        toc = doc.add_heading("", level=1)
        toc.add_run("فهرس المحتويات")
        _set_rtl(toc)
        
        for entry in _ALL_OUTPUTS:
            tp = doc.add_paragraph()
            tp.add_run(f"{entry['order']}. {entry['agent']} - {entry['time']}")
            _set_rtl(tp)
        
        doc.add_page_break()
        
        # All agent outputs
        main = doc.add_heading("", level=1)
        main.add_run("الحوار الكامل بين الوكلاء")
        _set_rtl(main)
        
        for entry in _ALL_OUTPUTS:
            # Agent header
            hdr = doc.add_paragraph()
            run = hdr.add_run(f"🤖 {entry['agent']}")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            _set_rtl(hdr)
            
            # Time
            tm = doc.add_paragraph()
            run = tm.add_run(f"الترتيب: {entry['order']} | الوقت: {entry['time']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            _set_rtl(tm)
            
            # Content
            for ln in entry["content"].splitlines():
                para = doc.add_paragraph()
                if ln.startswith("# "):
                    run = para.add_run(ln[2:])
                    run.bold = True
                    run.font.size = Pt(14)
                elif ln.startswith("## "):
                    run = para.add_run(ln[3:])
                    run.bold = True
                    run.font.size = Pt(12)
                elif ln.startswith("- ") or ln.startswith("* "):
                    para.add_run(f"• {ln[2:]}")
                elif ln.strip():
                    para.add_run(ln)
                _set_rtl(para)
            
            # Separator
            sep = doc.add_paragraph()
            sep.add_run("─" * 50)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
        
        # Save
        fpath = os.path.join(folder, "Rafeeq_Complete_Dialogue.docx")
        doc.save(fpath)
        
        return {
            "file_name": "Rafeeq_Complete_Dialogue.docx",
            "file_path": fpath,
            "status": "ok",
            "outputs_count": len(_ALL_OUTPUTS)
        }
    except Exception as err:
        return {"status": "error", "message": str(err)}
'''

def fix_all_tools(obj):
    changed = False
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool':
            lbl = obj.get('label', '')
            if 'export' in lbl.lower() or 'word' in lbl.lower():
                if 'config' in obj and 'source_code' in obj['config']:
                    print(f"  Fixing: {lbl}")
                    obj['config']['source_code'] = NEW_SOURCE_CODE
                    obj['config']['name'] = 'export_to_word'
                    changed = True
        for k, v in obj.items():
            if fix_all_tools(v):
                changed = True
    elif isinstance(obj, list):
        for itm in obj:
            if fix_all_tools(itm):
                changed = True
    return changed

def main():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    print("=== COMPREHENSIVE DIALOGUE EXPORT ===")
    
    cur.execute("SELECT id, component FROM team WHERE id = 4")
    row = cur.fetchone()
    if row:
        cfg = json.loads(row[1])
        if fix_all_tools(cfg):
            cur.execute("UPDATE team SET component = ? WHERE id = ?", 
                       (json.dumps(cfg), 4))
            print("  ✓ Rafeeq2 UPDATED")
    
    conn.commit()
    conn.close()
    print("\n✅ Export tool now captures ALL agent outputs!")
    print("   Each call accumulates content")
    print("   Final document contains complete dialogue")

if __name__ == "__main__":
    main()
