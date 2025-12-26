"""
Fix the Export To Word tool - simplified version without type hint issues
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Simplified export source code - avoiding complex type hints
NEW_SOURCE_CODE = '''from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Agent name mapping
AGENT_LABELS = {
    "rafeeq_moderator": ("Rafeeq Moderator", "وكيل المدير التنفيذي"),
    "series_bible_lock_agent": ("Series Bible Lock", "وكيل قفل الشخصيات"),
    "deep_research_agent": ("Deep Research", "وكيل البحث العميق"),
    "curriculum_engineer_agent": ("Curriculum Engineer", "مهندس المنهج"),
    "episode_planner_agent": ("Episode Planner", "وكيل تخطيط الحلقات"),
    "scenarist_agent": ("Scenarist v1", "كاتب السيناريو v1"),
    "sheikh_agent": ("Sheikh Reviewer", "المراجع الشرعي"),
    "education_agent": ("Education Consultant", "المستشار التربوي"),
    "psychiatrist_agent": ("Psychiatrist", "الطبيب النفسي"),
    "series_bible_agent": ("Series Bible Compliance", "وكيل اتساق السلسلة"),
    "youtube_consultant_agent": ("YouTube Consultant", "مستشار يوتيوب"),
    "language_consultant_agent": ("Language Consultant", "المدقق اللغوي"),
    "REVISION_DIGEST_AGENT": ("Revision Digest", "وكيل تلخيص المراجعات"),
    "scenarist_rewrite_agent": ("Scenarist v2", "كاتب السيناريو v2"),
    "COVERAGE_REPORT_AGENT": ("Coverage Report", "وكيل تقرير التغطية"),
    "rafeeq_exporter": ("Rafeeq Exporter", "وكيل التصدير")
}

# Global log
_LOG = []

def _set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

def _detect_agent(content):
    cl = content.lower()
    if "character_lock" in cl:
        return "series_bible_lock_agent"
    elif "child_friendly_definition" in cl:
        return "deep_research_agent"
    elif "learning_objective" in cl and "reinforcement" in cl:
        return "curriculum_engineer_agent"
    elif "beat_sheet" in cl:
        return "episode_planner_agent"
    elif "script_v1_title" in cl:
        return "scenarist_agent"
    elif "script_v2_text" in cl:
        return "scenarist_rewrite_agent"
    elif "script_text_final" in cl:
        return "language_consultant_agent"
    elif "patch_plan" in cl:
        return "REVISION_DIGEST_AGENT"
    elif "changelog" in cl:
        return "COVERAGE_REPORT_AGENT"
    elif "bible_compliance" in cl:
        return "series_bible_agent"
    elif "policy_and_safety" in cl:
        return "youtube_consultant_agent"
    elif "must_fix" in cl or "should_fix" in cl:
        if "شرع" in content:
            return "sheikh_agent"
        elif "تربو" in content:
            return "education_agent"
        elif "نفس" in content:
            return "psychiatrist_agent"
    return "rafeeq_moderator"

def export_to_word(report_markdown, section_title=None):
    """
    Export workflow dialogue to Word with Arabic RTL support.
    
    Args:
        report_markdown: The content to export
        section_title: Optional section title
    
    Returns:
        Dict with file info and status
    """
    global _LOG
    
    try:
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        os.makedirs(base_dir, exist_ok=True)
        
        file_name = "Rafeeq_Complete_Dialogue.docx"
        file_path = os.path.join(base_dir, file_name)
        
        # Detect agent
        agent_key = _detect_agent(report_markdown)
        agent_en, agent_ar = AGENT_LABELS.get(agent_key, (agent_key, agent_key))
        
        # Log
        _LOG.append({
            "agent_key": agent_key,
            "agent_en": agent_en,
            "agent_ar": agent_ar,
            "message": report_markdown,
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "order": len(_LOG) + 1
        })
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading("تقرير حوار وكلاء رفيق", level=0)
        _set_rtl(title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = doc.add_paragraph()
        info.add_run(f"تاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')} | عدد الرسائل: {len(_LOG)}")
        _set_rtl(info)
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents
        toc = doc.add_heading("فهرس الحوار", level=1)
        _set_rtl(toc)
        
        for entry in _LOG:
            p = doc.add_paragraph()
            p.add_run(f"{entry['order']}. {entry['agent_ar']} - {entry['timestamp']}")
            _set_rtl(p)
        
        doc.add_page_break()
        
        # Dialogue
        dh = doc.add_heading("الحوار الكامل", level=1)
        _set_rtl(dh)
        
        for entry in _LOG:
            # Agent header
            ap = doc.add_paragraph()
            run = ap.add_run(f"🤖 {entry['agent_ar']} ({entry['agent_en']})")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            _set_rtl(ap)
            
            # Timestamp
            ts = doc.add_paragraph()
            tsr = ts.add_run(f"⏱ {entry['timestamp']}")
            tsr.font.size = Pt(9)
            tsr.font.color.rgb = RGBColor(128, 128, 128)
            
            # Content
            for line in entry["message"].splitlines():
                mp = doc.add_paragraph()
                if line.startswith("# "):
                    r = mp.add_run(line[2:])
                    r.bold = True
                    r.font.size = Pt(14)
                elif line.startswith("## "):
                    r = mp.add_run(line[3:])
                    r.bold = True
                    r.font.size = Pt(12)
                elif line.startswith("- "):
                    mp.add_run(f"• {line[2:]}")
                else:
                    mp.add_run(line)
                _set_rtl(mp)
            
            # Separator
            sep = doc.add_paragraph("─" * 40)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
        
        doc.save(file_path)
        
        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
            "total_messages": len(_LOG),
            "current_agent": agent_ar
        }
    except Exception as e:
        return {"status": f"error:{str(e)}", "file_path": ""}
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                modified = True
                print(f"  ✓ Fixed function signature")
        
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== FIXING EXPORT TOOL FUNCTION SIGNATURE ===")
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"  ✓ Rafeeq2 team UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool fixed!")

if __name__ == "__main__":
    main()
