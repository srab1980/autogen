"""
Update rafeeq_exporter's export tool - fixed version with no short variable names
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Export tool with all long variable names - no 'para', 'p', etc.
AUTO_EXPORT_CODE = '''from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import sqlite3
import json
import os

AGENT_NAMES = {
    "rafeeq_moderator": ("Rafeeq Moderator", "وكيل المدير التنفيذي"),
    "series_bible_lock_agent": ("Series Bible Lock", "وكيل قفل الشخصيات"),
    "deep_research_agent": ("Deep Research", "وكيل البحث العميق"),
    "curriculum_engineer_agent": ("Curriculum Engineer", "مهندس المنهج"),
    "episode_planner_agent": ("Episode Planner", "وكيل تخطيط الحلقات"),
    "scenarist_agent": ("Scenarist v1", "كاتب السيناريو v1"),
    "sheikh_agent": ("Sheikh Reviewer", "المراجع الشرعي"),
    "education_agent": ("Education Consultant", "المستشار التربوي"),
    "psychiatrist_agent": ("Psychiatrist", "الطبيب النفسي"),
    "series_bible_agent": ("Series Bible Compliance", "وكيل اتساق السلسلة"),
    "youtube_consultant_agent": ("YouTube Consultant", "مستشار يوتيوب"),
    "language_consultant_agent": ("Language Consultant", "المدقق اللغوي"),
    "REVISION_DIGEST_AGENT": ("Revision Digest", "وكيل تلخيص المراجعات"),
    "scenarist_rewrite_agent": ("Scenarist v2", "كاتب السيناريو v2"),
    "COVERAGE_REPORT_AGENT": ("Coverage Report", "وكيل تقرير التغطية"),
    "rafeeq_exporter": ("Rafeeq Exporter", "وكيل التصدير"),
    "user": ("User", "المستخدم")
}

EXCLUDE_SOURCES = {"llm_call_event", "tool_call_event", "system", "orchestrator"}

def export_to_word(content: str, title: str = "") -> dict:
    """Export complete workflow dialogue to Word automatically."""
    try:
        db_path = "autogen04202.db"
        folder = "C:\\\\Users\\\\srab1.SAMEH-NVME\\\\Downloads\\\\AutoGen Studio Final\\\\AutoGen Studio\\\\AutoGen Studio\\\\Script"
        os.makedirs(folder, exist_ok=True)
        
        # Get latest run messages from database
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT r.id FROM run r
            JOIN session s ON r.session_id = s.id
            WHERE s.team_id = 4
            ORDER BY r.created_at DESC LIMIT 1
        """)
        run_row = cursor.fetchone()
        
        if not run_row:
            conn.close()
            return {"status": "error", "message": "No run found"}
        
        run_id = run_row[0]
        
        cursor.execute(
            "SELECT config FROM message WHERE run_id = ? ORDER BY created_at ASC",
            (run_id,)
        )
        
        messages = []
        for row in cursor.fetchall():
            if row[0]:
                try:
                    data = json.loads(row[0])
                    source = data.get("source", "")
                    msg_content = data.get("content", "")
                    
                    if source in EXCLUDE_SOURCES:
                        continue
                    if isinstance(msg_content, str) and msg_content.strip().startswith("{"):
                        continue
                    if not msg_content:
                        continue
                        
                    messages.append({"source": source, "content": msg_content})
                except:
                    pass
        
        conn.close()
        
        # Create Word document
        word_doc = Document()
        
        title_heading = word_doc.add_heading("تقرير حوار وكلاء رفيق الشامل", 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_paragraph = word_doc.add_paragraph()
        info_paragraph.add_run(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        count_paragraph = word_doc.add_paragraph()
        count_paragraph.add_run(f"عدد الرسائل: {len(messages)}")
        count_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        word_doc.add_page_break()
        
        # Table of contents
        word_doc.add_heading("فهرس المحتويات", 1)
        
        for idx, msg in enumerate(messages, 1):
            src = msg["source"]
            en_name, ar_name = AGENT_NAMES.get(src, (src, src))
            toc_para = word_doc.add_paragraph()
            toc_para.add_run(f"{idx}. {ar_name} ({en_name})")
        
        word_doc.add_page_break()
        
        # Full dialogue
        word_doc.add_heading("الحوار الكامل", 1)
        
        for idx, msg in enumerate(messages, 1):
            src = msg["source"]
            txt = msg["content"]
            en_name, ar_name = AGENT_NAMES.get(src, (src, src))
            
            agent_paragraph = word_doc.add_paragraph()
            agent_run = agent_paragraph.add_run(f"🤖 {ar_name} ({en_name})")
            agent_run.bold = True
            agent_run.font.size = Pt(14)
            agent_run.font.color.rgb = RGBColor(0, 102, 204)
            
            order_paragraph = word_doc.add_paragraph()
            order_run = order_paragraph.add_run(f"الترتيب: {idx}")
            order_run.font.size = Pt(9)
            order_run.font.color.rgb = RGBColor(128, 128, 128)
            
            if isinstance(txt, str):
                for text_line in txt.splitlines():
                    line_paragraph = word_doc.add_paragraph()
                    if text_line.startswith("# "):
                        line_run = line_paragraph.add_run(text_line[2:])
                        line_run.bold = True
                        line_run.font.size = Pt(14)
                    elif text_line.startswith("## "):
                        line_run = line_paragraph.add_run(text_line[3:])
                        line_run.bold = True
                        line_run.font.size = Pt(12)
                    elif text_line.startswith("- "):
                        line_paragraph.add_run(f"• {text_line[2:]}")
                    elif text_line.strip():
                        line_paragraph.add_run(text_line)
            
            separator_paragraph = word_doc.add_paragraph()
            separator_paragraph.add_run("─" * 50)
            separator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            word_doc.add_paragraph("")
        
        file_path = os.path.join(folder, "Rafeeq_Complete_Dialogue.docx")
        word_doc.save(file_path)
        
        return {"status": "ok", "file_path": file_path, "messages": len(messages)}
        
    except Exception as error_obj:
        return {"status": "error", "message": str(error_obj)}
'''

def update_exporter_tool():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== UPDATING EXPORTER TOOL (FIXED) ===")
    
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    config = json.loads(row[1])
    
    participants = config.get('config', {}).get('participants', [])
    
    for participant in participants:
        agent_name = participant.get('config', {}).get('name', '')
        
        if agent_name == 'rafeeq_exporter':
            workbenches = participant.get('config', {}).get('workbench', [])
            for wb in workbenches:
                tools = wb.get('config', {}).get('tools', [])
                for tool in tools:
                    if 'config' in tool:
                        tool['config']['source_code'] = AUTO_EXPORT_CODE
                        print(f"  ✓ Updated export tool on {agent_name}")
    
    cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                  (json.dumps(config), 4))
    conn.commit()
    conn.close()
    
    print("\n✅ Fixed! No more 'para' variable.")

if __name__ == "__main__":
    update_exporter_tool()
