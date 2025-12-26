"""
Create comprehensive workflow documentation export tool
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Comprehensive export source code that creates a single detailed report
NEW_SOURCE_CODE = '''from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Global variable to track workflow sections
_WORKFLOW_SECTIONS = []
_CURRENT_FILE = None

def export_to_word(report_markdown: str, section_title: Optional[str] = None, is_final: bool = False) -> Dict[str, str]:
    """
    Add a section to the comprehensive workflow report.
    All sections are collected and written to a single Word document.

    Parameters
    ----------
    report_markdown : str
        The content for this section.
    section_title : Optional[str]
        Title for this section (e.g., "Deep Research", "Script v1", "Reviews").
        If not provided, auto-detects from content.
    is_final : bool
        If True, this is the final export - writes everything to file.

    Returns
    -------
    Dict[str, str]
        Status information.
    """
    global _WORKFLOW_SECTIONS, _CURRENT_FILE
    
    try:
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        os.makedirs(base_dir, exist_ok=True)
        
        file_name = "Rafeeq_Workflow_Report.docx"
        file_path = os.path.join(base_dir, file_name)
        
        # Auto-detect section title from content if not provided
        if not section_title:
            content_lower = report_markdown.lower()
            if "series_bible" in content_lower or "character_lock" in content_lower:
                section_title = "1. Series Bible Lock"
            elif "research" in content_lower or "بحث" in report_markdown:
                section_title = "2. Deep Research"
            elif "learning_objective" in content_lower or "curriculum" in content_lower:
                section_title = "3. Curriculum Plan"
            elif "outline" in content_lower or "beat_sheet" in content_lower:
                section_title = "4. Episode Outline"
            elif "script_v1" in content_lower or "مسودة" in report_markdown:
                section_title = "5. Script v1 (First Draft)"
            elif "must_fix" in content_lower or "should_fix" in content_lower:
                section_title = "6. Reviewer Comments"
            elif "patch_plan" in content_lower:
                section_title = "7. Revision Plan"
            elif "script_v2" in content_lower:
                section_title = "8. Script v2 (Revised)"
            elif "script_text_final" in content_lower or "نهائي" in report_markdown:
                section_title = "9. Final Script"
            else:
                section_title = f"Section {len(_WORKFLOW_SECTIONS) + 1}"
        
        # Add section to our collection
        _WORKFLOW_SECTIONS.append({
            "title": section_title,
            "content": report_markdown,
            "timestamp": datetime.now().strftime("%H:%M:%S")
        })
        
        # Always write to file (update progressively)
        doc = Document()
        
        # Title Page
        title = doc.add_heading("تقرير سير عمل رفيق الشامل", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        subtitle = doc.add_paragraph("Rafeeq Workflow Documentation")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        doc.add_paragraph(f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"عدد المراحل المسجلة: {len(_WORKFLOW_SECTIONS)}")
        
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading("فهرس المحتويات", level=1)
        for i, section in enumerate(_WORKFLOW_SECTIONS, 1):
            doc.add_paragraph(f"{section['title']} - ({section['timestamp']})")
        
        doc.add_page_break()
        
        # Add all sections
        for section in _WORKFLOW_SECTIONS:
            # Section header
            doc.add_heading(section['title'], level=1)
            doc.add_paragraph(f"الوقت: {section['timestamp']}")
            doc.add_paragraph("─" * 50)
            
            # Section content
            for raw_line in section['content'].splitlines():
                line = raw_line.rstrip()
                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=2)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=3)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=4)
                elif line.startswith("- "):
                    p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.startswith("* "):
                    p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.strip() == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)
            
            doc.add_page_break()
        
        # Save
        doc.save(file_path)
        
        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
            "sections_count": len(_WORKFLOW_SECTIONS),
            "current_section": section_title
        }

    except Exception as e:
        return {
            "file_name": "",
            "file_path": "",
            "status": f"error:{str(e)}",
        }
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                obj['config']['description'] = "Comprehensive workflow documentation - exports all agent outputs to a single organized Word document"
                modified = True
                print(f"  ✓ Updated to comprehensive workflow mode")
        
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== UPDATING TO COMPREHENSIVE WORKFLOW EXPORT ===")
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"  ✓ Rafeeq2 team UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool now creates a COMPREHENSIVE workflow report:")
    print(f"   - Table of Contents")
    print(f"   - Each agent's output in order")  
    print(f"   - All reviewer comments")
    print(f"   - Final script")
    print(f"   All in ONE file: Rafeeq_Workflow_Report.docx")

if __name__ == "__main__":
    main()
