"""
Fix the Export To Word tool to save to the Script folder
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"
SCRIPT_FOLDER = r"C:\Users\srab1.SAMEH-NVME\Downloads\AutoGen Studio Final\AutoGen Studio\AutoGen Studio\Script"

# The corrected source code for the export function
NEW_SOURCE_CODE = '''from typing import Dict, Optional
from docx import Document
from datetime import datetime
import os


def export_to_word(report_markdown: str, file_name: Optional[str] = None) -> Dict[str, str]:
    """
    Create a .docx report file from a markdown-like text and return basic metadata.

    Parameters
    ----------
    report_markdown : str
        The full content of the report in simple markdown style.
        Lines starting with '# ' will be Heading 1,
        lines starting with '## ' will be Heading 2,
        all other lines become normal paragraphs.
    file_name : Optional[str]
        Optional base file name (without path). If not provided,
        a default name like 'Rafeeq_Report_YYYYMMDD_HHMMSS.docx' will be used.

    Returns
    -------
    Dict[str, str]
        A dictionary containing:
        - "file_name": the final file name
        - "file_path": the absolute path to the created file
        - "status": "ok" if created successfully, otherwise "error:<message>"
    """
    try:
        # Decide file name
        if file_name is None or not file_name.strip():
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"Rafeeq_Report_{timestamp}.docx"
        else:
            if not file_name.lower().endswith(".docx"):
                file_name = f"{file_name}.docx"

        # Use the Script folder as the output directory
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        
        # Create directory if it doesn't exist
        os.makedirs(base_dir, exist_ok=True)
        
        file_path = os.path.join(base_dir, file_name)

        # Create the document
        doc = Document()

        for raw_line in report_markdown.splitlines():
            line = raw_line.rstrip("\\n")
            if line.startswith("# "):
                # Heading 1
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                # Heading 2
                doc.add_heading(line[3:].strip(), level=2)
            else:
                # Normal paragraph (including empty lines)
                if line.strip() == "":
                    doc.add_paragraph("")  # blank line
                else:
                    doc.add_paragraph(line)

        doc.save(file_path)

        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
        }

    except Exception as e:
        return {
            "file_name": file_name or "",
            "file_path": "",
            "status": f"error:{str(e)}",
        }
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        # Check if this is a tool config with the export function
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                modified = True
                print(f"  ✓ Updated source_code to save to Script folder")
        
        # Recurse
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    # First, create the Script folder if it doesn't exist
    import os
    os.makedirs(SCRIPT_FOLDER, exist_ok=True)
    print(f"✓ Ensured Script folder exists: {SCRIPT_FOLDER}")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Fix Team 4 (Rafeeq2)
    print("\n=== FIXING EXPORT TOOL IN RAFEEQ2 TEAM ===")
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"  ✓ Team 4 UPDATED")
    
    # Also check Gallery for any export tools
    print("\n=== CHECKING GALLERY FOR EXPORT TOOLS ===")
    cursor.execute("SELECT id, config FROM gallery")
    for row_id, blob in cursor.fetchall():
        config = json.loads(blob)
        if fix_export_tool(config, f"gallery.{row_id}"):
            cursor.execute("UPDATE gallery SET config = ? WHERE id = ?", 
                         (json.dumps(config), row_id))
            print(f"  ✓ Gallery {row_id} UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool fixed! Files will now save to:\n   {SCRIPT_FOLDER}")

if __name__ == "__main__":
    main()
