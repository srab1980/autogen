"""
Export Complete Rafeeq Workflow Dialogue to Word
Reads messages from the 'message' table in AutoGen Studio's database.
Run this AFTER your Rafeeq workflow completes.
"""
import sqlite3
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Configuration
DB_PATH = "autogen04202.db"
OUTPUT_FOLDER = r"C:\Users\srab1.SAMEH-NVME\Downloads\AutoGen Studio Final\AutoGen Studio\AutoGen Studio\Script"
OUTPUT_FILE = "Rafeeq_Complete_Dialogue.docx"

# Agent name mapping
AGENT_DISPLAY_NAMES = {
    "rafeeq_moderator": ("Rafeeq Moderator", "وكيل المدير التنفيذي"),
    "series_bible_lock_agent": ("Series Bible Lock", "وكيل قفل الشخصيات"),
    "deep_research_agent": ("Deep Research", "وكيل البحث العميق"),
    "curriculum_engineer_agent": ("Curriculum Engineer", "مهندس المنهج"),
    "episode_planner_agent": ("Episode Planner", "وكيل تخطيط الحلقات"),
    "scenarist_agent": ("Scenarist v1", "كاتب السيناريو v1"),
    "sheikh_agent": ("Sheikh Reviewer", "المراجع الشرعي"),
    "education_agent": ("Education Consultant", "المستشار التربوي"),
    "psychiatrist_agent": ("Psychiatrist", "الطبيب النفسي"),
    "series_bible_agent": ("Series Bible Compliance", "وكيل اتساق السلسلة"),
    "youtube_consultant_agent": ("YouTube Consultant", "مستشار يوتيوب"),
    "language_consultant_agent": ("Language Consultant", "المدقق اللغوي"),
    "REVISION_DIGEST_AGENT": ("Revision Digest", "وكيل تلخيص المراجعات"),
    "scenarist_rewrite_agent": ("Scenarist v2", "كاتب السيناريو v2"),
    "COVERAGE_REPORT_AGENT": ("Coverage Report", "وكيل تقرير التغطية"),
    "rafeeq_exporter": ("Rafeeq Exporter", "وكيل التصدير"),
    "user": ("User", "المستخدم")
}

def set_rtl(paragraph):
    """Set RTL direction for Arabic text"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)

def get_latest_run_messages():
    """Get messages from the latest Rafeeq run"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get latest run for Rafeeq2 team
    cursor.execute("""
        SELECT r.id, r.created_at
        FROM run r
        JOIN session s ON r.session_id = s.id
        WHERE s.team_id = 4
        ORDER BY r.created_at DESC
        LIMIT 1
    """)
    
    run_row = cursor.fetchone()
    if not run_row:
        conn.close()
        return None, None, []
    
    run_id, created_at = run_row
    
    # Get messages for this run from the message table
    cursor.execute(
        "SELECT config, created_at FROM message WHERE run_id = ? ORDER BY created_at ASC",
        (run_id,)
    )
    
    # Sources to EXCLUDE (internal system messages)
    EXCLUDE_SOURCES = {
        "llm_call_event",
        "tool_call_event", 
        "system",
        "orchestrator"
    }
    
    # Message types to EXCLUDE
    EXCLUDE_TYPES = {
        "LLMCall",
        "ToolCall",
        "SystemMessage"
    }
    
    messages = []
    for row in cursor.fetchall():
        config_blob, msg_time = row
        if config_blob:
            try:
                msg_data = json.loads(config_blob)
                source = msg_data.get("source", "unknown")
                msg_type = msg_data.get("type", "")
                content = msg_data.get("content", "")
                
                # Skip internal/system messages
                if source in EXCLUDE_SOURCES:
                    continue
                if msg_type in EXCLUDE_TYPES:
                    continue
                # Skip if content looks like JSON (internal data)
                if isinstance(content, str) and content.strip().startswith('{"'):
                    continue
                # Skip empty content
                if not content or (isinstance(content, str) and not content.strip()):
                    continue
                    
                messages.append({
                    "source": source,
                    "content": content,
                    "type": msg_type,
                    "time": msg_time
                })
            except json.JSONDecodeError:
                pass
    
    conn.close()
    return run_id, created_at, messages

def export_to_word(run_id, created_at, messages):
    """Export messages to Word document"""
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    doc = Document()
    
    # Title page
    title = doc.add_heading("", level=0)
    title_run = title.add_run("تقرير حوار وكلاء رفيق الشامل")
    title_run.font.size = Pt(24)
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("Rafeeq Complete Workflow Dialogue")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.add_run(f"Run ID: {run_id}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_para.add_run(f"التاريخ: {created_at}")
    set_rtl(date_para)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    count_para = doc.add_paragraph()
    count_para.add_run(f"عدد الرسائل: {len(messages)}")
    set_rtl(count_para)
    count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of contents
    toc_title = doc.add_heading("فهرس المحتويات", level=1)
    set_rtl(toc_title)
    
    for i, msg in enumerate(messages, 1):
        source = msg.get("source", "unknown")
        en_name, ar_name = AGENT_DISPLAY_NAMES.get(source, (source, source))
        
        toc_item = doc.add_paragraph()
        toc_item.add_run(f"{i}. {ar_name} ({en_name})")
        set_rtl(toc_item)
    
    doc.add_page_break()
    
    # Full dialogue
    dialogue_title = doc.add_heading("الحوار الكامل بين الوكلاء", level=1)
    set_rtl(dialogue_title)
    
    for i, msg in enumerate(messages, 1):
        source = msg.get("source", "unknown")
        content = msg.get("content", "")
        msg_type = msg.get("type", "")
        
        en_name, ar_name = AGENT_DISPLAY_NAMES.get(source, (source, source))
        
        # Agent header
        agent_header = doc.add_paragraph()
        header_run = agent_header.add_run(f"🤖 {ar_name} ({en_name})")
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.color.rgb = RGBColor(0, 102, 204)
        set_rtl(agent_header)
        
        # Order and type
        order_para = doc.add_paragraph()
        order_run = order_para.add_run(f"الترتيب: {i} | النوع: {msg_type}")
        order_run.font.size = Pt(9)
        order_run.font.color.rgb = RGBColor(128, 128, 128)
        set_rtl(order_para)
        
        # Content
        if isinstance(content, str):
            for line in content.splitlines():
                content_para = doc.add_paragraph()
                
                if line.startswith("# "):
                    content_run = content_para.add_run(line[2:])
                    content_run.bold = True
                    content_run.font.size = Pt(14)
                elif line.startswith("## "):
                    content_run = content_para.add_run(line[3:])
                    content_run.bold = True
                    content_run.font.size = Pt(12)
                elif line.startswith("- ") or line.startswith("* "):
                    content_para.add_run(f"• {line[2:]}")
                elif line.strip():
                    content_para.add_run(line)
                
                set_rtl(content_para)
        else:
            content_para = doc.add_paragraph()
            content_para.add_run(str(content))
            set_rtl(content_para)
        
        # Separator
        separator = doc.add_paragraph()
        separator.add_run("─" * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    # Save
    output_path = os.path.join(OUTPUT_FOLDER, OUTPUT_FILE)
    doc.save(output_path)
    
    return output_path

def main():
    print("=" * 60)
    print("RAFEEQ COMPLETE DIALOGUE EXPORTER")
    print("=" * 60)
    
    print("\nFetching messages from latest Rafeeq run...")
    run_id, created_at, messages = get_latest_run_messages()
    
    if not run_id:
        print("❌ No Rafeeq runs found in database!")
        return
    
    print(f"✓ Found Run {run_id} with {len(messages)} messages")
    print(f"  Created: {created_at}")
    
    if len(messages) == 0:
        print("❌ No messages found for this run!")
        return
    
    # Show agents in this run
    print("\nAgents in this run:")
    sources = {}
    for msg in messages:
        source = msg.get("source", "unknown")
        sources[source] = sources.get(source, 0) + 1
    
    for source, count in sources.items():
        en_name, ar_name = AGENT_DISPLAY_NAMES.get(source, (source, source))
        print(f"  - {ar_name}: {count} messages")
    
    print("\nExporting to Word...")
    output_path = export_to_word(run_id, created_at, messages)
    
    print(f"\n✅ EXPORT COMPLETE!")
    print(f"   File: {output_path}")
    print(f"   Total Messages: {len(messages)}")
    print(f"   Unique Agents: {len(sources)}")

if __name__ == "__main__":
    main()
