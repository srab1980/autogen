"""
Update Export To Word tool with correct agent names
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Full agent list from user
AGENT_NAMES = {
    "rafeeq_moderator": "وكيل المدير التنفيذي",
    "series_bible_lock_agent": "وكيل قفل الشخصيات",
    "deep_research_agent": "وكيل البحث العميق",
    "curriculum_engineer_agent": "مهندس المنهج",
    "episode_planner_agent": "وكيل تخطيط الحلقات",
    "scenarist_agent": "كاتب السيناريو (v1)",
    "sheikh_agent": "المراجع الشرعي",
    "education_agent": "المستشار التربوي",
    "psychiatrist_agent": "الطبيب النفسي",
    "series_bible_agent": "وكيل اتساق السلسلة",
    "youtube_consultant_agent": "مستشار يوتيوب",
    "language_consultant_agent": "المدقق اللغوي",
    "REVISION_DIGEST_AGENT": "وكيل تلخيص المراجعات",
    "scenarist_rewrite_agent": "كاتب السيناريو (v2)",
    "COVERAGE_REPORT_AGENT": "وكيل تقرير التغطية",
    "rafeeq_exporter": "وكيل التصدير"
}

# Comprehensive export source code with all 16 agents
NEW_SOURCE_CODE = '''from typing import Dict, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Agent name mapping (English -> Arabic)
AGENT_LABELS = {
    "rafeeq_moderator": ("Rafeeq Moderator", "وكيل المدير التنفيذي"),
    "series_bible_lock_agent": ("Series Bible Lock", "وكيل قفل الشخصيات"),
    "deep_research_agent": ("Deep Research", "وكيل البحث العميق"),
    "curriculum_engineer_agent": ("Curriculum Engineer", "مهندس المنهج"),
    "episode_planner_agent": ("Episode Planner", "وكيل تخطيط الحلقات"),
    "scenarist_agent": ("Scenarist v1", "كاتب السيناريو v1"),
    "sheikh_agent": ("Sheikh Reviewer", "المراجع الشرعي"),
    "education_agent": ("Education Consultant", "المستشار التربوي"),
    "psychiatrist_agent": ("Psychiatrist", "الطبيب النفسي"),
    "series_bible_agent": ("Series Bible Compliance", "وكيل اتساق السلسلة"),
    "youtube_consultant_agent": ("YouTube Consultant", "مستشار يوتيوب"),
    "language_consultant_agent": ("Language Consultant", "المدقق اللغوي"),
    "REVISION_DIGEST_AGENT": ("Revision Digest", "وكيل تلخيص المراجعات"),
    "scenarist_rewrite_agent": ("Scenarist v2", "كاتب السيناريو v2"),
    "COVERAGE_REPORT_AGENT": ("Coverage Report", "وكيل تقرير التغطية"),
    "rafeeq_exporter": ("Rafeeq Exporter", "وكيل التصدير")
}

# Global conversation log
_CONVERSATION_LOG: List[Dict] = []

def _set_rtl(paragraph):
    """Set RTL direction for Arabic text"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def _detect_agent(content: str) -> str:
    """Auto-detect agent from content"""
    content_lower = content.lower()
    
    if "character_lock" in content_lower or "location_lock" in content_lower:
        return "series_bible_lock_agent"
    elif "child_friendly_definition" in content_lower or "story_seeds" in content_lower:
        return "deep_research_agent"
    elif "learning_objective" in content_lower and "reinforcement" in content_lower:
        return "curriculum_engineer_agent"
    elif "beat_sheet" in content_lower or "rafiq_interactive" in content_lower:
        return "episode_planner_agent"
    elif "script_v1_title" in content_lower or "script_v1_text" in content_lower:
        return "scenarist_agent"
    elif "script_v2_text" in content_lower or "coverage_section" in content_lower:
        return "scenarist_rewrite_agent"
    elif "script_text_final" in content_lower and "top_fixes" in content_lower:
        return "language_consultant_agent"
    elif "patch_plan" in content_lower:
        return "REVISION_DIGEST_AGENT"
    elif "changelog" in content_lower and "risk_check" in content_lower:
        return "COVERAGE_REPORT_AGENT"
    elif "export_status" in content_lower or "file_path" in content_lower:
        return "rafeeq_exporter"
    elif "bible_compliance" in content_lower:
        return "series_bible_agent"
    elif "policy_and_safety" in content_lower or "title_and_thumbnail" in content_lower:
        return "youtube_consultant_agent"
    elif ("must_fix" in content_lower or "should_fix" in content_lower):
        if "شرع" in content or "فقه" in content or "عقيد" in content:
            return "sheikh_agent"
        elif "تربو" in content or "إيجاب" in content:
            return "education_agent"
        elif "نفس" in content or "عاطف" in content:
            return "psychiatrist_agent"
        else:
            return "rafeeq_moderator"
    else:
        return "rafeeq_moderator"

def export_to_word(report_markdown: str, section_title: Optional[str] = None) -> Dict[str, str]:
    """
    Export comprehensive workflow dialogue to Word with Arabic RTL.
    Captures all 16 agents' outputs.
    """
    global _CONVERSATION_LOG
    
    try:
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        os.makedirs(base_dir, exist_ok=True)
        
        file_name = "Rafeeq_Complete_Dialogue.docx"
        file_path = os.path.join(base_dir, file_name)
        
        # Detect agent
        agent_key = _detect_agent(report_markdown)
        agent_en, agent_ar = AGENT_LABELS.get(agent_key, (agent_key, agent_key))
        
        # Log this message
        _CONVERSATION_LOG.append({
            "agent_key": agent_key,
            "agent_en": agent_en,
            "agent_ar": agent_ar,
            "message": report_markdown,
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "order": len(_CONVERSATION_LOG) + 1
        })
        
        # Create document
        doc = Document()
        
        # Title Page
        title = doc.add_heading("تقرير حوار وكلاء رفيق الشامل", level=0)
        _set_rtl(title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        p = doc.add_paragraph(f"تاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        _set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph(f"عدد الرسائل: {len(_CONVERSATION_LOG)}")
        _set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents
        toc = doc.add_heading("فهرس الحوار", level=1)
        _set_rtl(toc)
        toc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for entry in _CONVERSATION_LOG:
            p = doc.add_paragraph(f"{entry['order']}. {entry['agent_ar']} ({entry['agent_en']}) - {entry['timestamp']}")
            _set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_page_break()
        
        # Full Dialogue
        dialogue_header = doc.add_heading("الحوار الكامل", level=1)
        _set_rtl(dialogue_header)
        dialogue_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for entry in _CONVERSATION_LOG:
            # Agent header
            agent_p = doc.add_paragraph()
            _set_rtl(agent_p)
            run = agent_p.add_run(f"🤖 {entry['agent_ar']} ({entry['agent_en']})")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            agent_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Timestamp
            ts_p = doc.add_paragraph()
            run = ts_p.add_run(f"⏱ {entry['timestamp']} | الترتيب: {entry['order']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            _set_rtl(ts_p)
            ts_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Message content
            for line in entry['message'].splitlines():
                p = doc.add_paragraph()
                _set_rtl(p)
                
                if line.startswith("# "):
                    run = p.add_run(line[2:].strip())
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith("## "):
                    run = p.add_run(line[3:].strip())
                    run.bold = True
                    run.font.size = Pt(12)
                elif line.startswith("- ") or line.startswith("* "):
                    run = p.add_run(f"• {line[2:].strip()}")
                else:
                    run = p.add_run(line)
                
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Separator
            sep = doc.add_paragraph("─" * 50)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
        
        doc.save(file_path)
        
        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
            "total_messages": len(_CONVERSATION_LOG),
            "current_agent": agent_ar
        }

    except Exception as e:
        return {"status": f"error:{str(e)}"}
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                modified = True
                print(f"  ✓ Updated with all 16 agent names")
        
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== UPDATING WITH ALL 16 AGENT NAMES ===")
    print("Agents:")
    for key, name in AGENT_NAMES.items():
        print(f"  - {key}: {name}")
    
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"\n✓ Rafeeq2 team UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool now detects all 16 agents correctly!")

if __name__ == "__main__":
    main()
