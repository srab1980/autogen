"""
Create comprehensive dialogue export tool with Arabic RTL support
Captures ALL agent interactions in proper document format
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Comprehensive export source code with Arabic RTL support
NEW_SOURCE_CODE = '''from typing import Dict, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Global conversation log
_CONVERSATION_LOG: List[Dict] = []

def _set_rtl(paragraph):
    """Set RTL direction for Arabic text"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def _set_document_rtl(doc):
    """Set document-level RTL settings"""
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)

def _add_agent_header(doc, agent_name: str, role: str = ""):
    """Add a styled agent header"""
    p = doc.add_paragraph()
    _set_rtl(p)
    
    run = p.add_run(f"🤖 {agent_name}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    if role:
        role_run = p.add_run(f" ({role})")
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = RGBColor(128, 128, 128)
    
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _add_message(doc, content: str, indent_level: int = 0):
    """Add a message with proper Arabic RTL formatting"""
    for line in content.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
            
        p = doc.add_paragraph()
        _set_rtl(p)
        
        # Handle markdown-like formatting
        if line.startswith("# "):
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("## "):
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith("### "):
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("- ") or line.startswith("* "):
            run = p.add_run(f"• {line[2:].strip()}")
            run.font.size = Pt(11)
        else:
            run = p.add_run(line)
            run.font.size = Pt(11)
        
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add indentation
        if indent_level > 0:
            p.paragraph_format.right_indent = Inches(0.3 * indent_level)

def log_agent_message(agent_name: str, message: str, role: str = "") -> Dict[str, str]:
    """
    Log an agent's message for later export.
    Call this for EVERY agent output to capture the full dialogue.
    
    Parameters
    ----------
    agent_name : str
        Name of the agent (e.g., "Deep Research Agent", "Sheikh")
    message : str
        The agent's full output/response
    role : str
        Optional role description
        
    Returns
    -------
    Dict with status
    """
    global _CONVERSATION_LOG
    
    _CONVERSATION_LOG.append({
        "agent": agent_name,
        "role": role,
        "message": message,
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "order": len(_CONVERSATION_LOG) + 1
    })
    
    return {
        "status": "logged",
        "agent": agent_name,
        "message_count": len(_CONVERSATION_LOG)
    }

def export_to_word(report_markdown: str, section_title: Optional[str] = None) -> Dict[str, str]:
    """
    Export the complete workflow dialogue to a Word document.
    Captures all agent interactions with proper Arabic RTL formatting.

    Parameters
    ----------
    report_markdown : str
        Content to add (or use "EXPORT_ALL" to finalize)
    section_title : Optional[str]
        Title for this section

    Returns
    -------
    Dict with file info
    """
    global _CONVERSATION_LOG
    
    try:
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        os.makedirs(base_dir, exist_ok=True)
        
        file_name = "Rafeeq_Complete_Dialogue.docx"
        file_path = os.path.join(base_dir, file_name)
        
        # Log this call as well
        if section_title:
            agent_name = section_title
        else:
            # Auto-detect agent from content
            content_lower = report_markdown.lower()
            if "series_bible" in content_lower or "character_lock" in content_lower:
                agent_name = "Series Bible Lock Agent"
            elif "research" in content_lower or "بحث" in report_markdown:
                agent_name = "Deep Research Agent"
            elif "curriculum" in content_lower or "learning_objective" in content_lower:
                agent_name = "Curriculum Engineer Agent"
            elif "outline" in content_lower or "beat_sheet" in content_lower:
                agent_name = "Episode Planner Agent"
            elif "script_v1" in content_lower:
                agent_name = "Scenarist Agent (v1)"
            elif "must_fix" in content_lower or "should_fix" in content_lower:
                if "شرع" in report_markdown:
                    agent_name = "Sheikh (Sharia Reviewer)"
                elif "تربو" in report_markdown or "education" in content_lower:
                    agent_name = "Education Consultant"
                elif "نفس" in report_markdown or "psychiatr" in content_lower:
                    agent_name = "Psychiatrist"
                elif "youtube" in content_lower:
                    agent_name = "YouTube Consultant"
                elif "bible" in content_lower:
                    agent_name = "Series Bible Agent"
                else:
                    agent_name = "Reviewer Agent"
            elif "patch_plan" in content_lower:
                agent_name = "Revision Digest Agent"
            elif "script_v2" in content_lower:
                agent_name = "Scenarist Rewrite Agent"
            elif "script_text_final" in content_lower or "top_fixes" in content_lower:
                agent_name = "Language Consultant"
            elif "export" in content_lower:
                agent_name = "Rafeeq Exporter"
            else:
                agent_name = "Rafeeq Moderator"
        
        log_agent_message(agent_name, report_markdown)
        
        # Create comprehensive document
        doc = Document()
        
        # Set RTL for document
        _set_document_rtl(doc)
        
        # Title Page
        title = doc.add_heading("", level=0)
        run = title.add_run("تقرير حوار وكلاء رفيق الشامل")
        run.font.size = Pt(24)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rtl(title)
        
        doc.add_paragraph("")
        
        subtitle = doc.add_paragraph()
        run = subtitle.add_run("Rafeeq Agents Complete Dialogue")
        run.font.size = Pt(16)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        
        info = doc.add_paragraph()
        _set_rtl(info)
        run = info.add_run(f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info2 = doc.add_paragraph()
        _set_rtl(info2)
        run = info2.add_run(f"عدد الرسائل المسجلة: {len(_CONVERSATION_LOG)}")
        info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents
        toc_title = doc.add_heading("", level=1)
        run = toc_title.add_run("فهرس الحوار")
        _set_rtl(toc_title)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for entry in _CONVERSATION_LOG:
            p = doc.add_paragraph()
            _set_rtl(p)
            run = p.add_run(f"{entry['order']}. {entry['agent']} - ({entry['timestamp']})")
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_page_break()
        
        # Full Dialogue
        dialogue_title = doc.add_heading("", level=1)
        run = dialogue_title.add_run("الحوار الكامل بين الوكلاء")
        _set_rtl(dialogue_title)
        dialogue_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("")
        
        for entry in _CONVERSATION_LOG:
            # Agent header
            _add_agent_header(doc, entry['agent'], entry.get('role', ''))
            
            # Timestamp
            ts = doc.add_paragraph()
            _set_rtl(ts)
            run = ts.add_run(f"⏱ {entry['timestamp']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            ts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Message content
            _add_message(doc, entry['message'])
            
            # Separator
            sep = doc.add_paragraph()
            run = sep.add_run("─" * 60)
            run.font.color.rgb = RGBColor(200, 200, 200)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("")
        
        doc.save(file_path)
        
        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
            "total_messages": len(_CONVERSATION_LOG),
            "agents_logged": list(set(e['agent'] for e in _CONVERSATION_LOG))
        }

    except Exception as e:
        return {
            "file_name": "",
            "file_path": "",
            "status": f"error:{str(e)}",
        }

def reset_conversation_log() -> Dict[str, str]:
    """Reset the conversation log for a new workflow run"""
    global _CONVERSATION_LOG
    _CONVERSATION_LOG = []
    return {"status": "reset", "message_count": 0}
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                obj['config']['description'] = "Export complete agent dialogue with Arabic RTL support"
                modified = True
                print(f"  ✓ Updated with Arabic RTL support")
        
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== UPDATING TO COMPREHENSIVE DIALOGUE EXPORT WITH ARABIC RTL ===")
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"  ✓ Rafeeq2 team UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool now creates:")
    print(f"   - Complete dialogue between ALL agents")
    print(f"   - Proper Arabic RTL formatting")
    print(f"   - Structured document with Table of Contents")
    print(f"   - Agent identification for each message")
    print(f"   - Timestamps for each exchange")
    print(f"\n   Output: Script\\Rafeeq_Complete_Dialogue.docx")

if __name__ == "__main__":
    main()
