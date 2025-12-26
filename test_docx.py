"""
Test the export tool to see if it creates valid Word files
"""
from docx import Document
from datetime import datetime
import os

# Test creating a simple document
folder = os.path.join(
    "C:", os.sep, "Users", "srab1.SAMEH-NVME", "Downloads",
    "AutoGen Studio Final", "AutoGen Studio", "AutoGen Studio", "Script"
)
os.makedirs(folder, exist_ok=True)

fpath = os.path.join(folder, "Test_Document.docx")

doc = Document()
doc.add_heading("Test Document", level=0)
doc.add_paragraph(f"Created: {datetime.now()}")
doc.add_paragraph("")
doc.add_heading("This is a test", level=1)
doc.add_paragraph("If you can read this, the document was created successfully!")
doc.add_paragraph("")
doc.add_paragraph("تجربة النص العربي: هذا نص تجريبي")

doc.save(fpath)
print(f"✅ Test document created: {fpath}")
print(f"   File size: {os.path.getsize(fpath)} bytes")
