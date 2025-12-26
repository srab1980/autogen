"""
Fixed comprehensive export - using explicit variable names throughout
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Fixed export tool with explicit variable names (no para, ln, etc.)
NEW_SOURCE_CODE = '''from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

SAVED_OUTPUTS = []
CALL_COUNT = 0

def set_arabic_rtl(paragraph_obj):
    paragraph_pr = paragraph_obj._p.get_or_add_pPr()
    bidi_elem = OxmlElement("w:bidi")
    bidi_elem.set(qn("w:val"), "1")
    paragraph_pr.append(bidi_elem)

AGENT_NAMES = {
    "bible_lock": "وكيل قفل الشخصيات",
    "research": "وكيل البحث العميق",
    "curriculum": "مهندس المنهج",
    "planner": "وكيل تخطيط الحلقات",
    "script1": "كاتب السيناريو v1",
    "script2": "كاتب السيناريو v2",
    "sheikh": "المراجع الشرعي",
    "education": "المستشار التربوي",
    "psychiatrist": "الطبيب النفسي",
    "bible_check": "وكيل اتساق السلسلة",
    "youtube": "مستشار يوتيوب",
    "revision": "وكيل تلخيص المراجعات",
    "coverage": "وكيل تقرير التغطية",
    "language": "المدقق اللغوي",
    "exporter": "وكيل التصدير",
    "moderator": "وكيل المدير التنفيذي"
}

def detect_agent_name(text_content: str) -> str:
    lower_text = text_content.lower()
    if "character_lock" in lower_text:
        return AGENT_NAMES["bible_lock"]
    elif "child_friendly_definition" in lower_text:
        return AGENT_NAMES["research"]
    elif "learning_objective" in lower_text and "key_takeaways" in lower_text:
        return AGENT_NAMES["curriculum"]
    elif "beat_sheet" in lower_text:
        return AGENT_NAMES["planner"]
    elif "script_v1_title" in lower_text:
        return AGENT_NAMES["script1"]
    elif "script_v2_title" in lower_text:
        return AGENT_NAMES["script2"]
    elif "script_text_final" in lower_text:
        return AGENT_NAMES["language"]
    elif "patch_plan" in lower_text:
        return AGENT_NAMES["revision"]
    elif "changelog" in lower_text:
        return AGENT_NAMES["coverage"]
    elif "bible_compliance" in lower_text:
        return AGENT_NAMES["bible_check"]
    elif "policy_and_safety" in lower_text:
        return AGENT_NAMES["youtube"]
    elif "must_fix" in lower_text:
        if "شرع" in text_content:
            return AGENT_NAMES["sheikh"]
        elif "تربو" in text_content:
            return AGENT_NAMES["education"]
        elif "نفس" in text_content:
            return AGENT_NAMES["psychiatrist"]
    elif "export_status" in lower_text:
        return AGENT_NAMES["exporter"]
    return AGENT_NAMES["moderator"]

def export_to_word(content: str, title: str = "") -> dict:
    """Export all agent outputs to comprehensive Word document."""
    global SAVED_OUTPUTS, CALL_COUNT
    
    try:
        output_folder = "C:\\\\Users\\\\srab1.SAMEH-NVME\\\\Downloads\\\\AutoGen Studio Final\\\\AutoGen Studio\\\\AutoGen Studio\\\\Script"
        os.makedirs(output_folder, exist_ok=True)
        
        agent_display_name = detect_agent_name(str(content))
        
        CALL_COUNT = CALL_COUNT + 1
        SAVED_OUTPUTS.append({
            "order_num": CALL_COUNT,
            "agent_name": agent_display_name,
            "text_content": str(content),
            "time_stamp": datetime.now().strftime("%H:%M:%S")
        })
        
        word_document = Document()
        
        title_heading = word_document.add_heading("تقرير حوار وكلاء رفيق", level=0)
        set_arabic_rtl(title_heading)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_paragraph = word_document.add_paragraph()
        date_paragraph.add_run(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        set_arabic_rtl(date_paragraph)
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        count_paragraph = word_document.add_paragraph()
        count_paragraph.add_run(f"عدد المخرجات: {len(SAVED_OUTPUTS)}")
        set_arabic_rtl(count_paragraph)
        count_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        word_document.add_page_break()
        
        toc_heading = word_document.add_heading("فهرس المحتويات", level=1)
        set_arabic_rtl(toc_heading)
        
        for saved_entry in SAVED_OUTPUTS:
            toc_item = word_document.add_paragraph()
            toc_item.add_run(f"{saved_entry['order_num']}. {saved_entry['agent_name']} - {saved_entry['time_stamp']}")
            set_arabic_rtl(toc_item)
        
        word_document.add_page_break()
        
        dialogue_heading = word_document.add_heading("الحوار الكامل", level=1)
        set_arabic_rtl(dialogue_heading)
        
        for saved_entry in SAVED_OUTPUTS:
            agent_header = word_document.add_paragraph()
            header_run = agent_header.add_run(f"🤖 {saved_entry['agent_name']}")
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.color.rgb = RGBColor(0, 102, 204)
            set_arabic_rtl(agent_header)
            
            time_info = word_document.add_paragraph()
            time_run = time_info.add_run(f"الترتيب: {saved_entry['order_num']} | الوقت: {saved_entry['time_stamp']}")
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = RGBColor(128, 128, 128)
            set_arabic_rtl(time_info)
            
            for text_line in saved_entry["text_content"].splitlines():
                content_paragraph = word_document.add_paragraph()
                if text_line.startswith("# "):
                    line_run = content_paragraph.add_run(text_line[2:])
                    line_run.bold = True
                    line_run.font.size = Pt(14)
                elif text_line.startswith("## "):
                    line_run = content_paragraph.add_run(text_line[3:])
                    line_run.bold = True
                    line_run.font.size = Pt(12)
                elif text_line.startswith("- "):
                    content_paragraph.add_run(f"• {text_line[2:]}")
                elif text_line.strip():
                    content_paragraph.add_run(text_line)
                set_arabic_rtl(content_paragraph)
            
            separator_paragraph = word_document.add_paragraph()
            separator_paragraph.add_run("─" * 50)
            separator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            word_document.add_paragraph("")
        
        file_path = os.path.join(output_folder, "Rafeeq_Complete_Dialogue.docx")
        word_document.save(file_path)
        
        return {"file_name": "Rafeeq_Complete_Dialogue.docx", "file_path": file_path, "status": "ok", "count": len(SAVED_OUTPUTS)}
    except Exception as error_obj:
        return {"status": "error", "message": str(error_obj)}
'''

def fix_tools(config_obj):
    modified = False
    if isinstance(config_obj, dict):
        if config_obj.get('component_type') == 'tool':
            label_text = config_obj.get('label', '')
            if 'export' in label_text.lower() or 'word' in label_text.lower():
                if 'config' in config_obj and 'source_code' in config_obj['config']:
                    print(f"  Fixing: {label_text}")
                    config_obj['config']['source_code'] = NEW_SOURCE_CODE
                    config_obj['config']['name'] = 'export_to_word'
                    modified = True
        for key_name, value_obj in config_obj.items():
            if fix_tools(value_obj):
                modified = True
    elif isinstance(config_obj, list):
        for item_obj in config_obj:
            if fix_tools(item_obj):
                modified = True
    return modified

def main():
    db_conn = sqlite3.connect(DB_PATH)
    db_cursor = db_conn.cursor()
    
    print("=== FIXING WITH EXPLICIT VARIABLE NAMES ===")
    
    db_cursor.execute("SELECT id, component FROM team WHERE id = 4")
    team_row = db_cursor.fetchone()
    if team_row:
        team_config = json.loads(team_row[1])
        if fix_tools(team_config):
            db_cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                             (json.dumps(team_config), 4))
            print("  ✓ Rafeeq2 UPDATED")
    
    db_conn.commit()
    db_conn.close()
    print("\n✅ Fixed with explicit variable names!")

if __name__ == "__main__":
    main()
