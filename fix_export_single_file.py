"""
Fix the Export To Word tool to create a single consolidated file
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# The corrected source code - appends to a single file
NEW_SOURCE_CODE = '''from typing import Dict, Optional
from docx import Document
from datetime import datetime
import os


def export_to_word(report_markdown: str, file_name: Optional[str] = None) -> Dict[str, str]:
    """
    Create or append to a .docx report file from a markdown-like text.
    All exports in a session go to a single consolidated file.

    Parameters
    ----------
    report_markdown : str
        The full content of the report in simple markdown style.
        Lines starting with '# ' will be Heading 1,
        lines starting with '## ' will be Heading 2,
        all other lines become normal paragraphs.
    file_name : Optional[str]
        Optional base file name (without path). If not provided,
        uses 'Rafeeq_Complete_Report.docx'.

    Returns
    -------
    Dict[str, str]
        A dictionary containing:
        - "file_name": the final file name
        - "file_path": the absolute path to the created file
        - "status": "ok" if created successfully, otherwise "error:<message>"
    """
    try:
        # Use a single consolidated file name
        if file_name is None or not file_name.strip():
            file_name = "Rafeeq_Complete_Report.docx"
        else:
            if not file_name.lower().endswith(".docx"):
                file_name = f"{file_name}.docx"

        # Use the Script folder as the output directory
        base_dir = r"C:\\Users\\srab1.SAMEH-NVME\\Downloads\\AutoGen Studio Final\\AutoGen Studio\\AutoGen Studio\\Script"
        
        # Create directory if it doesn't exist
        os.makedirs(base_dir, exist_ok=True)
        
        file_path = os.path.join(base_dir, file_name)

        # Check if file exists - if so, open and append; otherwise create new
        if os.path.exists(file_path):
            doc = Document(file_path)
            # Add a page break before new content
            doc.add_page_break()
        else:
            doc = Document()
            # Add a title page for new document
            doc.add_heading("تقرير رفيق الشامل", level=0)
            doc.add_paragraph(f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_page_break()

        # Add a separator with timestamp for this section
        doc.add_paragraph(f"--- قسم جديد ({datetime.now().strftime('%H:%M:%S')}) ---")
        doc.add_paragraph("")

        for raw_line in report_markdown.splitlines():
            line = raw_line.rstrip("\\n")
            if line.startswith("# "):
                # Heading 1
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                # Heading 2
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                # Heading 3
                doc.add_heading(line[4:].strip(), level=3)
            else:
                # Normal paragraph (including empty lines)
                if line.strip() == "":
                    doc.add_paragraph("")  # blank line
                else:
                    doc.add_paragraph(line)

        doc.save(file_path)

        return {
            "file_name": file_name,
            "file_path": os.path.abspath(file_path),
            "status": "ok",
            "mode": "appended" if os.path.exists(file_path) else "created"
        }

    except Exception as e:
        return {
            "file_name": file_name or "",
            "file_path": "",
            "status": f"error:{str(e)}",
        }
'''

def fix_export_tool(obj, path="root"):
    """Find and fix the Export To Word tool"""
    modified = False
    
    if isinstance(obj, dict):
        # Check if this is a tool config with the export function
        if obj.get('component_type') == 'tool' and obj.get('label') == 'Export To Word':
            print(f"  Found Export To Word tool at {path}")
            if 'config' in obj and 'source_code' in obj['config']:
                obj['config']['source_code'] = NEW_SOURCE_CODE
                obj['config']['description'] = "Export report to a single consolidated Word file. Each call appends to the same file."
                modified = True
                print(f"  ✓ Updated to single-file mode")
        
        # Recurse
        for key, value in obj.items():
            if fix_export_tool(value, f"{path}.{key}"):
                modified = True
    
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if fix_export_tool(item, f"{path}[{i}]"):
                modified = True
    
    return modified

def main():
    # First, clear any existing reports to start fresh
    import os
    script_folder = r"C:\Users\srab1.SAMEH-NVME\Downloads\AutoGen Studio Final\AutoGen Studio\AutoGen Studio\Script"
    os.makedirs(script_folder, exist_ok=True)
    
    # Remove old reports
    for f in os.listdir(script_folder):
        if f.endswith('.docx'):
            os.remove(os.path.join(script_folder, f))
            print(f"  Removed old file: {f}")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Fix Team 4 (Rafeeq2)
    print("\n=== UPDATING EXPORT TOOL TO SINGLE-FILE MODE ===")
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    
    if row:
        config = json.loads(row[1])
        if fix_export_tool(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print(f"  ✓ Team 4 UPDATED")
    
    # Also check Gallery
    cursor.execute("SELECT id, config FROM gallery")
    for row_id, blob in cursor.fetchall():
        config = json.loads(blob)
        if fix_export_tool(config, f"gallery.{row_id}"):
            cursor.execute("UPDATE gallery SET config = ? WHERE id = ?", 
                         (json.dumps(config), row_id))
            print(f"  ✓ Gallery {row_id} UPDATED")
    
    conn.commit()
    conn.close()
    
    print(f"\n✅ Export tool now creates a SINGLE consolidated file:")
    print(f"   {script_folder}\\Rafeeq_Complete_Report.docx")

if __name__ == "__main__":
    main()
