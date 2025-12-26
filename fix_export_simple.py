"""
Simple Export To Word tool - no global state, no complex features
"""
import sqlite3
import json

DB_PATH = "autogen04202.db"

# Simple, working export tool
NEW_SOURCE_CODE = '''from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

def export_to_word(report_markdown, section_title=None):
    """Export content to Word file with Arabic support."""
    try:
        # Output directory
        base_dir = os.path.join(
            "C:", os.sep, "Users", "srab1.SAMEH-NVME", "Downloads",
            "AutoGen Studio Final", "AutoGen Studio", "AutoGen Studio", "Script"
        )
        os.makedirs(base_dir, exist_ok=True)
        
        file_name = "Rafeeq_Report.docx"
        file_path = os.path.join(base_dir, file_name)
        
        # Open existing or create new
        if os.path.exists(file_path):
            doc = Document(file_path)
            doc.add_page_break()
        else:
            doc = Document()
            doc.add_heading("تقرير رفيق", level=0)
            doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
            doc.add_page_break()
        
        # Add section header
        if section_title:
            doc.add_heading(section_title, level=1)
        else:
            doc.add_heading(f"Section ({datetime.now().strftime('%H:%M:%S')})", level=1)
        
        # Add content
        for line in str(report_markdown).splitlines():
            if line.startswith("# "):
                p = doc.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        
        doc.add_paragraph("")
        doc.add_paragraph("-" * 40)
        
        doc.save(file_path)
        
        return {
            "file_name": file_name,
            "file_path": file_path,
            "status": "ok"
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}
'''

def fix_all_export_tools(obj):
    """Find and fix ALL Export To Word tools in config"""
    modified = False
    
    if isinstance(obj, dict):
        if obj.get('component_type') == 'tool':
            label = obj.get('label', '')
            if 'export' in label.lower() or 'word' in label.lower():
                if 'config' in obj and 'source_code' in obj['config']:
                    print(f"  Fixing tool: {label}")
                    obj['config']['source_code'] = NEW_SOURCE_CODE
                    obj['config']['name'] = 'export_to_word'
                    obj['config']['description'] = 'Export report to Word file'
                    modified = True
        
        for key, value in obj.items():
            if fix_all_export_tools(value):
                modified = True
    
    elif isinstance(obj, list):
        for item in obj:
            if fix_all_export_tools(item):
                modified = True
    
    return modified

def main():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("=== APPLYING SIMPLE EXPORT TOOL ===")
    
    # Fix Rafeeq2 team
    cursor.execute("SELECT id, component FROM team WHERE id = 4")
    row = cursor.fetchone()
    if row:
        config = json.loads(row[1])
        if fix_all_export_tools(config):
            cursor.execute("UPDATE team SET component = ? WHERE id = ?", 
                         (json.dumps(config), 4))
            print("  ✓ Rafeeq2 UPDATED")
    
    conn.commit()
    conn.close()
    print("\n✅ Simple export tool applied!")

if __name__ == "__main__":
    main()
